"""
HTML-to-DOCX converter using python-docx + BeautifulSoup.

Replaces the pdf2docx pipeline (HTML→PDF→DOCX) with direct HTML→DOCX conversion.
This preserves text semantics, spacing, bilingual structure, and images correctly
instead of relying on position-based PDF text extraction which merges sentences.
"""

import base64
import io
import logging
import re
from typing import Optional

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Tags that should create block-level breaks
BLOCK_TAGS = {
    "div", "p", "h1", "h2", "h3", "h4", "h5", "h6",
    "ul", "ol", "li", "table", "blockquote", "section",
    "article", "header", "footer", "figure", "figcaption",
}

# Tags that are inline formatting
INLINE_TAGS = {"b", "strong", "i", "em", "u", "span", "sup", "sub", "a", "br", "img"}

# Heading tag to point size mapping
HEADING_SIZES = {
    "h1": 17,
    "h2": 15,
    "h3": 14,
    "h4": 13,
    "h5": 12,
    "h6": 11,
}


class HtmlToDocxConverter:
    """
    Converts reconstructed HTML pages into a well-formatted DOCX document.

    Handles:
    - Block elements (div, p, headings)
    - Inline formatting (bold, italic, color, superscript)
    - Images (base64 data URIs → embedded pictures)
    - Tables (with proper cell formatting)
    - Bilingual text (English + native language stacked via <br>)
    - Math expressions (rendered as text; MathJax SVG not available in DOCX)
    - MCQ option layouts (flex/grid → sequential paragraphs)
    """

    def __init__(self, lang_code: str = ""):
        self.lang_code = lang_code
        self.doc = Document()
        self._setup_document_styles()

    def _setup_document_styles(self):
        """Configure default document styles for A4 layout."""
        # Set A4 page size and margins
        for section in self.doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        # Set default font
        style = self.doc.styles["Normal"]
        font = style.font
        font.size = Pt(11)
        font.name = "Calibri"

        # Set default paragraph spacing
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.15

    def convert(self, pages: dict) -> io.BytesIO:
        """
        Convert all completed pages into a DOCX document.

        Args:
            pages: dict of page_number -> page_data with reconstructed_html

        Returns:
            BytesIO buffer containing the .docx file
        """
        sorted_pages = sorted(pages.items(), key=lambda x: int(x[0]))
        total = len(sorted_pages)

        for idx, (page_num, page_data) in enumerate(sorted_pages):
            if page_data.get("status") != "completed":
                continue

            html = page_data.get("reconstructed_html", "")
            if not html:
                continue

            logger.info(f"DOCX: Processing page {page_num} ({len(html)} chars HTML)")

            try:
                self._process_page_html(html, int(page_num))
            except Exception as e:
                logger.error(f"DOCX: Failed to process page {page_num}: {e}", exc_info=True)
                # Add error placeholder
                p = self.doc.add_paragraph()
                run = p.add_run(f"[Page {page_num} - rendering error]")
                run.font.color.rgb = RGBColor(200, 0, 0)
                run.font.size = Pt(10)

            # Add a subtle page separator (except after last page)
            if idx < total - 1:
                self._add_page_separator()

        # Save to buffer
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        logger.info(f"DOCX: Document generated, {len(sorted_pages)} pages, {buffer.getbuffer().nbytes:,} bytes")
        return buffer

    def _process_page_html(self, html: str, page_number: int):
        """Parse a single page's HTML and add elements to the document."""
        soup = BeautifulSoup(html, "html.parser")

        # Find the main content container
        # The HTML is wrapped in <div class="translated-page">
        main_div = soup.find("div", class_="translated-page")
        if not main_div:
            # No wrapper — process the soup directly
            main_div = soup

        # Strip <style> tags (not needed in DOCX)
        for style_tag in main_div.find_all("style"):
            style_tag.decompose()

        # Process all children
        self._process_children(main_div)

    def _process_children(self, parent: Tag):
        """Recursively process child elements of a parent tag."""
        for child in parent.children:
            if isinstance(child, NavigableString):
                text = str(child)
                # Skip pure whitespace between block elements
                if text.strip():
                    p = self.doc.add_paragraph()
                    self._add_styled_text(p, text.strip())
            elif isinstance(child, Tag):
                self._process_element(child)

    def _process_element(self, element: Tag):
        """Process a single HTML element and add to document."""
        tag = element.name

        if tag is None:
            return

        # Skip invisible elements
        style = element.get("style", "")
        if "display:none" in style.replace(" ", "") or "display: none" in style:
            return

        # --- Images ---
        if tag == "img":
            self._handle_image(element)
            return

        # --- Tables ---
        if tag == "table":
            self._handle_table(element)
            return

        # --- Headings ---
        if tag in HEADING_SIZES:
            self._handle_heading(element)
            return

        # --- Lists ---
        if tag in ("ul", "ol"):
            self._handle_list(element, ordered=(tag == "ol"))
            return

        # --- Line break at top level ---
        if tag == "br":
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            return

        # --- Horizontal rule (we skip per original design) ---
        if tag == "hr":
            return

        # --- Block-level divs and paragraphs ---
        if tag in ("div", "p", "section", "article", "blockquote"):
            self._handle_block(element)
            return

        # --- Inline elements at root level (shouldn't normally happen) ---
        if tag in INLINE_TAGS:
            p = self.doc.add_paragraph()
            self._render_inline(p, element)
            return

        # --- Fallback: try processing children ---
        self._process_children(element)

    def _handle_block(self, element: Tag):
        """
        Handle a block-level element (div, p, section, etc.).
        
        Checks if the block contains sub-blocks (nested divs/tables/images)
        or is a leaf text block. For leaf blocks, renders as a single paragraph.
        For container blocks, recurses into children.
        """
        # Check if this block has only inline content (text, spans, br, b, etc.)
        has_block_children = False
        for child in element.children:
            if isinstance(child, Tag) and child.name in (BLOCK_TAGS | {"img", "table"}):
                has_block_children = True
                break

        if has_block_children:
            # This is a container div — recurse
            self._process_children(element)
        else:
            # This is a leaf text block — render as one paragraph
            text_content = element.get_text(strip=True)
            if not text_content:
                return

            p = self.doc.add_paragraph()

            # Check for special block styling
            style = element.get("style", "")
            if "text-align:center" in style.replace(" ", "") or "text-align: center" in style:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if "text-align:right" in style.replace(" ", "") or "text-align: right" in style:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Check for background color (header bars)
            bg_match = re.search(r'background-color\s*:\s*([^;"]+)', style)
            if bg_match:
                self._set_paragraph_shading(p, bg_match.group(1).strip())

            # Render inline content
            self._render_inline(p, element)

    def _render_inline(self, paragraph, element: Tag):
        """
        Render an element's inline content into an existing paragraph.
        Handles nested bold, italic, spans with color, <br>, <sup>, etc.
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    # Collapse excessive whitespace but preserve single spaces
                    text = re.sub(r'\s+', ' ', text)
                    if text.strip() or text == ' ':
                        self._add_styled_text(paragraph, text)

            elif isinstance(child, Tag):
                tag = child.name

                if tag == "br":
                    # Add a line break within the same paragraph
                    run = paragraph.add_run()
                    run.add_break()
                    continue

                if tag == "img":
                    # Inline image
                    self._handle_image(child, paragraph)
                    continue

                if tag in ("b", "strong"):
                    self._render_inline_with_format(paragraph, child, bold=True)
                    continue

                if tag in ("i", "em"):
                    self._render_inline_with_format(paragraph, child, italic=True)
                    continue

                if tag == "u":
                    self._render_inline_with_format(paragraph, child, underline=True)
                    continue

                if tag == "sup":
                    self._render_inline_with_format(paragraph, child, superscript=True)
                    continue

                if tag == "sub":
                    self._render_inline_with_format(paragraph, child, subscript=True)
                    continue

                if tag == "span":
                    # Check for color styling
                    span_style = child.get("style", "")
                    color = self._extract_color(span_style)
                    self._render_inline_with_format(paragraph, child, color=color)
                    continue

                if tag == "a":
                    # Render link text (no hyperlink in simple mode)
                    self._render_inline(paragraph, child)
                    continue

                # For any other inline tag, just recurse
                if tag in BLOCK_TAGS:
                    # Nested block inside what we thought was inline — handle as separate block
                    # This can happen with nested divs for MCQ options
                    self._handle_block(child)
                else:
                    self._render_inline(paragraph, child)

    def _render_inline_with_format(
        self, paragraph, element: Tag,
        bold=False, italic=False, underline=False,
        superscript=False, subscript=False,
        color: Optional[RGBColor] = None,
    ):
        """Render inline content with specific formatting applied."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    text = re.sub(r'\s+', ' ', text)
                    if text.strip() or text == ' ':
                        text = self._sanitize_text(text)
                        if not text:
                            continue
                        run = paragraph.add_run(text)
                        if bold:
                            run.bold = True
                        if italic:
                            run.italic = True
                        if underline:
                            run.underline = True
                        if superscript:
                            run.font.superscript = True
                        if subscript:
                            run.font.subscript = True
                        if color:
                            run.font.color.rgb = color
            elif isinstance(child, Tag):
                if child.name == "br":
                    run = paragraph.add_run()
                    run.add_break()
                elif child.name in ("b", "strong"):
                    self._render_inline_with_format(
                        paragraph, child,
                        bold=True, italic=italic, underline=underline,
                        superscript=superscript, subscript=subscript, color=color,
                    )
                elif child.name in ("i", "em"):
                    self._render_inline_with_format(
                        paragraph, child,
                        bold=bold, italic=True, underline=underline,
                        superscript=superscript, subscript=subscript, color=color,
                    )
                elif child.name == "sup":
                    self._render_inline_with_format(
                        paragraph, child,
                        bold=bold, italic=italic, underline=underline,
                        superscript=True, subscript=False, color=color,
                    )
                elif child.name == "sub":
                    self._render_inline_with_format(
                        paragraph, child,
                        bold=bold, italic=italic, underline=underline,
                        superscript=False, subscript=True, color=color,
                    )
                elif child.name == "span":
                    span_color = self._extract_color(child.get("style", "")) or color
                    self._render_inline_with_format(
                        paragraph, child,
                        bold=bold, italic=italic, underline=underline,
                        superscript=superscript, subscript=subscript, color=span_color,
                    )
                elif child.name == "img":
                    self._handle_image(child, paragraph)
                else:
                    self._render_inline_with_format(
                        paragraph, child,
                        bold=bold, italic=italic, underline=underline,
                        superscript=superscript, subscript=subscript, color=color,
                    )

    def _add_styled_text(self, paragraph, text: str):
        """Add plain text as a run to a paragraph."""
        if not text:
            return
        text = self._sanitize_text(text)
        if not text:
            return
        run = paragraph.add_run(text)
        run.font.size = Pt(11)

    @staticmethod
    def _sanitize_text(text: str) -> str:
        """Remove XML-incompatible control characters from text.
        
        python-docx / lxml rejects NULL bytes and most C0 control characters.
        We keep tab (0x09), newline (0x0A), and carriage return (0x0D).
        """
        if not text:
            return text
        # Remove all control chars except \t \n \r
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    def _handle_heading(self, element: Tag):
        """Add a heading paragraph."""
        tag = element.name
        level_map = {"h1": 0, "h2": 1, "h3": 2, "h4": 3, "h5": 4, "h6": 5}
        level = level_map.get(tag, 2)

        # python-docx heading levels: 0=Title, 1=Heading1, etc.
        heading_level = max(1, min(level + 1, 4))

        text = element.get_text(strip=True)
        if not text:
            return

        heading = self.doc.add_heading(text, level=heading_level)
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(3)

    def _handle_image(self, element: Tag, paragraph=None):
        """
        Handle an <img> tag. If src is a base64 data URI, embed it.
        Otherwise skip (broken image).
        """
        src = element.get("src", "")

        if not src.startswith("data:image"):
            # Skip non-data-URI images (broken/placeholder)
            return

        try:
            # Parse data URI: data:image/png;base64,XXXXX
            header, b64_data = src.split(",", 1)
            img_bytes = base64.b64decode(b64_data)

            # Determine image dimensions and scale
            from PIL import Image as PILImage
            pil_img = PILImage.open(io.BytesIO(img_bytes))
            img_w, img_h = pil_img.size

            # Calculate width: max 15cm (A4 content width with margins)
            max_width_cm = 15.0
            # Scale based on image aspect ratio
            aspect = img_h / img_w if img_w > 0 else 1
            width_cm = min(max_width_cm, img_w / 72 * 2.54)  # rough pixels to cm
            if width_cm < 3:
                width_cm = 3  # minimum 3cm

            width = Cm(width_cm)

            if paragraph is None:
                # Create a new centered paragraph for the image
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = paragraph

            run = p.add_run()
            run.add_picture(io.BytesIO(img_bytes), width=width)

        except Exception as e:
            logger.warning(f"DOCX: Failed to embed image: {e}")

    def _handle_table(self, element: Tag):
        """Convert an HTML table to a DOCX table."""
        rows_data = []
        max_cols = 0

        # Parse all rows
        for tr in element.find_all("tr", recursive=False):
            cells = tr.find_all(["td", "th"], recursive=False)
            row = []
            for cell in cells:
                cell_text = cell.get_text(separator="\n", strip=True)
                is_header = cell.name == "th"
                row.append({"text": cell_text, "is_header": is_header, "element": cell})
            rows_data.append(row)
            max_cols = max(max_cols, len(row))

        # Handle tables inside thead/tbody
        if not rows_data:
            for section in element.find_all(["thead", "tbody", "tfoot"], recursive=False):
                for tr in section.find_all("tr", recursive=False):
                    cells = tr.find_all(["td", "th"], recursive=False)
                    row = []
                    for cell in cells:
                        cell_text = cell.get_text(separator="\n", strip=True)
                        is_header = cell.name == "th"
                        row.append({"text": cell_text, "is_header": is_header, "element": cell})
                    rows_data.append(row)
                    max_cols = max(max_cols, len(row))

        if not rows_data or max_cols == 0:
            return

        # Create table
        table = self.doc.add_table(rows=len(rows_data), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style: use Table Grid for visible borders
        try:
            table.style = self.doc.styles["Table Grid"]
        except KeyError:
            pass

        # Fill cells
        for r_idx, row in enumerate(rows_data):
            for c_idx, cell_data in enumerate(row):
                if c_idx >= max_cols:
                    break
                cell = table.cell(r_idx, c_idx)
                # Clear default paragraph
                cell.paragraphs[0].clear()

                # Render cell content with formatting
                self._render_inline(cell.paragraphs[0], cell_data["element"])

                # Header formatting
                if cell_data["is_header"]:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                    # Header background
                    self._set_cell_shading(cell, "F0F0F0")

                # Cell font size
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)

        # Add spacing after table
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def _handle_list(self, element: Tag, ordered: bool = False, level: int = 0):
        """Handle <ul> and <ol> lists."""
        items = element.find_all("li", recursive=False)
        for i, li in enumerate(items):
            # Check for nested lists
            nested_list = li.find(["ul", "ol"], recursive=False)

            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 * (level + 1))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            # Add bullet/number prefix
            if ordered:
                prefix = f"{i + 1}. "
            else:
                prefix = "• "
            run = p.add_run(prefix)
            run.font.size = Pt(11)

            # Add list item text (excluding nested lists)
            for child in li.children:
                if isinstance(child, Tag) and child.name in ("ul", "ol"):
                    continue  # Handle nested list separately
                elif isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        run = p.add_run(text)
                        run.font.size = Pt(11)
                elif isinstance(child, Tag):
                    self._render_inline(p, child)

            # Handle nested lists
            if nested_list:
                self._handle_list(
                    nested_list,
                    ordered=(nested_list.name == "ol"),
                    level=level + 1,
                )

    def _add_page_separator(self):
        """Add a subtle visual separator between pages."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        # Thin gray line
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ─── Utility Methods ──────────────────────────────────────────

    def _extract_color(self, style: str) -> Optional[RGBColor]:
        """Extract color from inline CSS style string."""
        if not style:
            return None

        # Match color: #RRGGBB or color: #RGB
        hex_match = re.search(r'color\s*:\s*#([0-9a-fA-F]{3,6})', style)
        if hex_match:
            hex_val = hex_match.group(1)
            if len(hex_val) == 3:
                hex_val = ''.join(c * 2 for c in hex_val)
            try:
                r = int(hex_val[0:2], 16)
                g = int(hex_val[2:4], 16)
                b = int(hex_val[4:6], 16)
                return RGBColor(r, g, b)
            except (ValueError, IndexError):
                return None

        # Match color: rgb(r, g, b)
        rgb_match = re.search(r'color\s*:\s*rgb\((\d+),\s*(\d+),\s*(\d+)\)', style)
        if rgb_match:
            try:
                r = int(rgb_match.group(1))
                g = int(rgb_match.group(2))
                b = int(rgb_match.group(3))
                return RGBColor(r, g, b)
            except (ValueError, IndexError):
                return None

        return None

    def _set_paragraph_shading(self, paragraph, color_str: str):
        """Set background color on a paragraph."""
        hex_color = self._css_color_to_hex(color_str)
        if not hex_color:
            return

        pPr = paragraph._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)

    def _set_cell_shading(self, cell, hex_color: str):
        """Set background color on a table cell."""
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def _css_color_to_hex(self, color_str: str) -> Optional[str]:
        """Convert a CSS color value to a 6-char hex string."""
        color_str = color_str.strip().lower()

        # Already hex
        if color_str.startswith("#"):
            hex_val = color_str[1:]
            if len(hex_val) == 3:
                hex_val = ''.join(c * 2 for c in hex_val)
            return hex_val.upper() if len(hex_val) == 6 else None

        # rgb()
        rgb_match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_str)
        if rgb_match:
            r, g, b = int(rgb_match.group(1)), int(rgb_match.group(2)), int(rgb_match.group(3))
            return f"{r:02X}{g:02X}{b:02X}"

        # Named colors (common ones)
        named = {
            "white": "FFFFFF", "black": "000000", "red": "FF0000",
            "blue": "0000FF", "green": "008000", "gray": "808080",
            "grey": "808080", "lightgray": "D3D3D3", "lightgrey": "D3D3D3",
            "yellow": "FFFF00", "orange": "FFA500",
        }
        return named.get(color_str)
